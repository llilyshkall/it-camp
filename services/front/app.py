from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import os
import io
from datetime import datetime

import pandas as pd  # обработка Excel
from docx import Document  # формирование DOCX
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook  # для объединения Excel файлов
from openpyxl.utils.dataframe import dataframe_to_rows

app = Flask(__name__)

# ---------- Настройки загрузки ----------
app.config['MAX_CONTENT_LENGTH'] = 25 * 1024 * 1024  # 25 МБ суммарно на запрос
ALLOWED_EXTS = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.csv', '.txt'}
EXCEL_EXTS = {'.xls', '.xlsx'}


def allowed(filename: str) -> bool:
    _, ext = os.path.splitext((filename or '').lower())
    return ext in ALLOWED_EXTS


# ---------- Проверка «подходит для ашуренса» (заглушка-эвристика) ----------
def assurance_check(files) -> dict:
    reasons = []
    if not files:
        return {'ok': False, 'reasons': ['Файлы не получены']}

    total = 0
    ok_type = False
    keywords = ('реестр', 'портфель', 'договор', 'кредит', 'задолж', 'акт', 'оценк')

    for f in files:
        fname = secure_filename(f.filename or '')
        blob = f.read() or b''
        total += len(blob)
        f.seek(0)

        if allowed(fname):
            ok_type = True

        # Простейшая текстовая эвристика для txt/csv
        name, ext = os.path.splitext(fname.lower())
        if ext in ('.txt', '.csv'):
            try:
                text = (f.read() or b'').decode('utf-8', errors='ignore')
                f.seek(0)
                if any(k in text.lower() for k in keywords):
                    reasons.append(f'Найдены ключевые слова в «{fname}»')
            except Exception:
                pass

    if not ok_type:
        reasons.append('Нет файлов допустимых типов')
    if total > app.config['MAX_CONTENT_LENGTH']:
        reasons.append('Суммарный размер превышает лимит 25 МБ')

    ok = ok_type and total <= app.config['MAX_CONTENT_LENGTH']
    if not reasons:
        reasons.append('Базовая проверка пройдена')

    return {'ok': ok, 'reasons': reasons}


# ---------- Обработка замечаний (Excel) ----------
def process_remarks_excel(file_storage) -> bytes:
    """
    Принимает Excel с замечаниями и возвращает обработанный Excel (bytes).
    Здесь базовая обработка — замени под свои правила по мере необходимости.
    """
    fname = secure_filename(file_storage.filename or '')
    _, ext = os.path.splitext(fname.lower())
    if ext not in EXCEL_EXTS:
        raise ValueError("Ожидается файл Excel (.xls/.xlsx)")

    df = pd.read_excel(file_storage)

    # нормализация имён
    df.columns = (
        df.columns
        .map(lambda c: str(c).strip().lower())
        .map(lambda c: c.replace(' ', '_'))
    )

    # трим строковые поля
    for col in df.columns:
        if pd.api.types.is_string_dtype(df[col]):
            df[col] = df[col].astype(str).str.strip()

    # удаление дублей по ID + замечанию (если есть)
    subset = [c for c in ('id', 'замечание') if c in df.columns]
    if subset:
        df = df.drop_duplicates(subset=subset, keep='first')

    # пример статуса
    if 'статус' in df.columns:
        df['статус'] = df['статус'].fillna('').astype(str).str.strip()
        df['статус_норм'] = df['статус'].replace({'': 'К обработке'})
    else:
        df['статус_норм'] = 'К обработке'

    # сохраняем в память
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='замечания_обработанные')
    bio.seek(0)
    return bio.read()


# ---------- Формирование протокола (DOCX) ----------
def build_protocol_docx() -> bytes:
    """
    Генерирует простой протокол в формате DOCX.
    При желании можно прокинуть параметры (номер, стороны, даты, итоги) из формы.
    """
    doc = Document()

    # Титул
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('ПРОТОКОЛ\nпроверки и обработки материалов')
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph()  # пустая строка

    # Шапка
    now = datetime.now().strftime('%d.%m.%Y %H:%M')
    p = doc.add_paragraph()
    p.add_run('Дата и время формирования: ').bold = True
    p.add_run(now)

    # Разделы (скелет)
    doc.add_heading('1. Основание', level=2)
    doc.add_paragraph('Указать основание, номер задания/договора, ссылки на документы.')

    doc.add_heading('2. Результаты предварительной проверки («ашуренс»)', level=2)
    doc.add_paragraph('Указать итог: Подходит / Не подходит для ашуренса, краткие пояснения.')

    doc.add_heading('3. Обработка замечаний эксперта', level=2)
    doc.add_paragraph('Кратко описать, сколько замечаний получено, сколько обработано, статус оставшихся.')

    doc.add_heading('4. Решение / Рекомендации', level=2)
    doc.add_paragraph('Сформулировать итоговые выводы и последующие шаги.')

    # Подписи
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    table.cell(0, 0).text = 'Ответственный (ФИО/должность):'
    table.cell(0, 1).text = 'Подпись:'
    table.cell(1, 0).text = 'Согласовано:'
    table.cell(1, 1).text = 'Подпись:'

    # Вывод в память
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


# ---------- Маршруты ----------
@app.get("/")
def index():
    return render_template("index.html", page_title="Проект для Газпрома")


@app.get("/dashboards")
def dashboards():
    return render_template("dashboards.html")


@app.get("/pumpjack-demo")
def pumpjack_demo():
    return render_template("pumpjack-demo.html")


@app.get("/pumpjack-integration")
def pumpjack_integration():
    return render_template("pumpjack-integration.html")


@app.post("/assurance")
def assurance():
    files = request.files.getlist('files[]')
    result = assurance_check(files)
    resp = {
        'verdict': 'ok' if result['ok'] else 'fail',
        'title': 'Подходит для ашуренса' if result['ok'] else 'Не подходит для ашуренса',
        'reasons': result['reasons'],
    }
    return jsonify(resp), 200


@app.post("/remarks")
def remarks():
    """
    Принимает несколько файлов Excel (input name='remarks[]'), отдаёт обработанный Excel как attachment.
    """
    files = request.files.getlist('remarks[]')
    if not files or not any(f.filename for f in files):
        return jsonify({'error': 'Файлы не переданы'}), 400

    try:
        # Обрабатываем все файлы и объединяем результаты
        all_data = []
        processed_files = []
        
        for file in files:
            if file and file.filename:
                try:
                    data = process_remarks_excel(file)
                    all_data.append(data)
                    processed_files.append(file.filename)
                except Exception as e:
                    # Логируем ошибку для конкретного файла, но продолжаем обработку остальных
                    print(f"Ошибка обработки файла {file.filename}: {e}")
                    continue
        
        if not all_data:
            return jsonify({'error': 'Не удалось обработать ни один файл'}), 400
        
        # Если обработан только один файл, возвращаем его как есть
        if len(all_data) == 1:
            return send_file(
                io.BytesIO(all_data[0]),
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name='remarks_processed.xlsx'
            )
        
        # Если несколько файлов, объединяем их в один Excel с несколькими листами
        from openpyxl import Workbook
        from openpyxl.utils.dataframe import dataframe_to_rows
        
        wb = Workbook()
        # Удаляем дефолтный лист
        wb.remove(wb.active)
        
        for i, (data, filename) in enumerate(zip(all_data, processed_files)):
            try:
                # Читаем данные из bytes
                df = pd.read_excel(io.BytesIO(data))
                
                # Создаем лист для каждого файла
                sheet_name = f"Файл_{i+1}_{os.path.splitext(filename)[0][:20]}"
                ws = wb.create_sheet(title=sheet_name)
                
                # Записываем данные в лист
                for row in dataframe_to_rows(df, index=False, header=True):
                    ws.append(row)
                    
            except Exception as e:
                print(f"Ошибка создания листа для {filename}: {e}")
                continue
        
        # Сохраняем объединенный файл
        bio = io.BytesIO()
        wb.save(bio)
        bio.seek(0)
        
        return send_file(
            bio,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'remarks_processed_{len(processed_files)}_files.xlsx'
        )
        
    except Exception as e:
        return jsonify({'error': f'Не удалось обработать файлы: {e}'}), 400


@app.post("/protocol")
def protocol():
    """
    Формирует протокол (DOCX) и отдаёт на скачивание.
    При необходимости позже примем параметры из формы.
    """
    data = build_protocol_docx()
    return send_file(
        io.BytesIO(data),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='protocol.docx'
    )


# В dev-режиме можно запускать напрямую
if __name__ == "__main__":
    app.run(debug=True)
